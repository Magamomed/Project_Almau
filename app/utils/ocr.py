import os
import fitz  # PyMuPDF
import io
from docx import Document
import easyocr
from PIL import Image

reader = easyocr.Reader(['ru', 'en'])

def extract_text_from_image(image_path):
    result = reader.readtext(image_path, detail=0, paragraph=True)
    return "\n".join(result)

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    full_text = ""

    for page in doc:
        pix = page.get_pixmap(dpi=300)
        img_bytes = pix.tobytes("png")

        image = Image.open(io.BytesIO(img_bytes))
        image_path = "temp_image.png"
        image.save(image_path)

        result = reader.readtext(image_path, detail=0, paragraph=True)
        full_text += "\n".join(result) + "\n"

        os.remove(image_path)

    return full_text

def extract_text_from_docx(docx_path):
    document = Document(docx_path)
    full_text = []

    # Сначала обычные параграфы
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    # Затем таблицы — построчно
    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                full_text.append(" | ".join(row_text))  # Разделим содержимое ячеек через |
    
    return "\n".join(full_text)
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext in ['.png', '.jpg', '.jpeg']:
        return extract_text_from_image(file_path)
    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        return "Неподдерживаемый формат файла"
